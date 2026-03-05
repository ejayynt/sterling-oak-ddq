import json
import io
import os
import tempfile

from contextlib import asynccontextmanager
from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, Form
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from docx import Document

from backend.database import init_db, get_db, GeneratedAnswer, User
from backend.rag import ingest_reference_document, generate_answer_for_question
from backend.auth import router as auth_router, get_current_user


# ---------------------------------------------------------------------------
# App lifecycle — initialise DB on startup
# ---------------------------------------------------------------------------
@asynccontextmanager
async def lifespan(app: FastAPI):
    init_db()
    yield

app = FastAPI(title="Sterling Oak DDQ Automator", lifespan=lifespan)

# Include the auth router for /signup + /login
app.include_router(auth_router)

# Allow frontend to communicate with backend (fallback for dev)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Temporary storage for the original uploaded questionnaire
TEMP_DIR = tempfile.gettempdir()


# ---------------------------------------------------------------------------
# Upload reference documents
# ---------------------------------------------------------------------------
@app.post("/upload_references")
async def upload_references(
    files: list[UploadFile] = File(...),
    current_user: User = Depends(get_current_user),
):
    for file in files:
        file_path = os.path.join(TEMP_DIR, file.filename)
        with open(file_path, "wb") as f:
            f.write(await file.read())
        ingest_reference_document(file_path, file.filename, current_user.id)

    return {"message": f"Successfully ingested {len(files)} reference document(s)."}


# ---------------------------------------------------------------------------
# Process questionnaire — parse + generate answers
# ---------------------------------------------------------------------------
@app.post("/process_questionnaire")
async def process_questionnaire(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Save the original so we can modify it on export
    original_path = os.path.join(TEMP_DIR, file.filename)
    with open(original_path, "wb") as f:
        f.write(await file.read())

    doc = Document(original_path)

    if not doc.tables:
        raise HTTPException(status_code=400, detail="No table found in the document.")

    table = doc.tables[0]
    results = []

    # Clear previous answers for this file so re-processing is idempotent
    db.query(GeneratedAnswer).filter(
        GeneratedAnswer.user_id == current_user.id,
        GeneratedAnswer.source_filename == file.filename,
    ).delete()
    db.commit()

    for i, row in enumerate(table.rows):
        # Skip the header row
        if i == 0 and "question" in row.cells[0].text.lower():
            continue

        question_text = row.cells[0].text.strip()
        if not question_text:
            continue

        try:
            ai_response = generate_answer_for_question(question_text, current_user.id)
        except Exception as exc:
            import logging
            logging.getLogger(__name__).error(f"AI generation failed for question: {question_text[:80]}... Error: {exc}")
            ai_response = {
                "answer": "Not found in references.",
                "citation": "None",
                "snippet": "None",
                "confidence": 0,
            }

        db_answer = GeneratedAnswer(
            user_id=current_user.id,
            question_text=question_text,
            answer_text=ai_response["answer"],
            citation=ai_response["citation"],
            evidence_snippet=ai_response["snippet"],
            confidence=ai_response["confidence"],
            source_filename=file.filename,
        )
        db.add(db_answer)
        db.commit()
        db.refresh(db_answer)

        results.append({
            "id": db_answer.id,
            "question": question_text,
            "answer": ai_response["answer"],
            "citation": ai_response["citation"],
            "snippet": ai_response["snippet"],
            "confidence": ai_response["confidence"],
        })

    return {"filename": file.filename, "data": results}


# ---------------------------------------------------------------------------
# Get saved answers (for refreshing the review view)
# ---------------------------------------------------------------------------
@app.get("/answers")
def get_answers(
    filename: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    rows = (
        db.query(GeneratedAnswer)
        .filter(
            GeneratedAnswer.user_id == current_user.id,
            GeneratedAnswer.source_filename == filename,
        )
        .order_by(GeneratedAnswer.id)
        .all()
    )
    return {
        "filename": filename,
        "data": [
            {
                "id": r.id,
                "question": r.question_text,
                "answer": r.answer_text,
                "citation": r.citation,
                "snippet": r.evidence_snippet,
                "confidence": r.confidence,
            }
            for r in rows
        ],
    }


# ---------------------------------------------------------------------------
# Partial regeneration — re-run RAG for a single answer
# ---------------------------------------------------------------------------
@app.post("/regenerate/{answer_id}")
def regenerate_answer(
    answer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    row = db.query(GeneratedAnswer).filter(GeneratedAnswer.id == answer_id).first()
    if not row or row.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Answer not found")

    ai = generate_answer_for_question(row.question_text, current_user.id)
    row.answer_text = ai["answer"]
    row.citation = ai["citation"]
    row.evidence_snippet = ai["snippet"]
    row.confidence = ai["confidence"]
    db.commit()
    db.refresh(row)

    return {
        "id": row.id,
        "question": row.question_text,
        "answer": row.answer_text,
        "citation": row.citation,
        "snippet": row.evidence_snippet,
        "confidence": row.confidence,
    }


# ---------------------------------------------------------------------------
# Edit a single answer (save edits from the review UI)
# ---------------------------------------------------------------------------
@app.put("/answers/{answer_id}")
def update_answer(
    answer_id: int,
    new_text: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    row = db.query(GeneratedAnswer).filter(GeneratedAnswer.id == answer_id).first()
    if not row or row.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Answer not found")
    row.answer_text = new_text
    db.commit()
    return {"status": "saved"}


# ---------------------------------------------------------------------------
# Export final document
# ---------------------------------------------------------------------------
@app.post("/export")
async def export_document(
    filename: str = Form(...),
    edited_answers: str = Form(...),
    current_user: User = Depends(get_current_user),
):
    answers_dict = json.loads(edited_answers)

    original_path = os.path.join(TEMP_DIR, filename)
    if not os.path.exists(original_path):
        raise HTTPException(status_code=404, detail="Original questionnaire file not found. Please re-upload.")

    doc = Document(original_path)
    table = doc.tables[0]

    for row in table.rows:
        question_text = row.cells[0].text.strip()
        if question_text in answers_dict:
            entry = answers_dict[question_text]
            final_text = f"{entry['answer']}\n\n[Source: {entry['citation']}]"
            row.cells[1].text = final_text

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Completed_{filename}"'},
    )


# ---------------------------------------------------------------------------
# Serve frontend — mount AFTER all API routes
# ---------------------------------------------------------------------------
FRONTEND_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "frontend")
FRONTEND_DIR = os.path.normpath(FRONTEND_DIR)
app.mount("/", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")
